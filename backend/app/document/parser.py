"""
文档解析器

支持格式：
- .txt / .md / .markdown：纯文本
- .pdf：使用 PyMuPDF（优先）或 pdfplumber 或 pypdf
- .docx：使用 python-docx
"""

import io
import tempfile
from typing import Optional, Dict, Any
from pathlib import Path
from loguru import logger


class DocumentParser:
    """文档解析器"""

    SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx"}

    def __init__(self, pdf_parser: str = "pymupdf"):
        """
        Args:
            pdf_parser: PDF 解析器选择 ("pymupdf", "pdfplumber", "pypdf")
        """
        self.pdf_parser = pdf_parser

    def parse(self, content: bytes, filename: str) -> str:
        """
        解析文档内容为纯文本

        Args:
            content: 文件二进制内容
            filename: 文件名（用于判断格式）

        Returns:
            str: 解析后的纯文本
        """
        ext = Path(filename).suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {ext}，支持: {self.SUPPORTED_EXTENSIONS}")

        if ext in (".txt", ".md", ".markdown"):
            return self._parse_text(content)
        if ext == ".pdf":
            return self._parse_pdf(content, filename)
        if ext == ".docx":
            return self._parse_docx_enhanced(content)

        return ""

    def parse_pdf_detailed(self, content: bytes, filename: str) -> Dict[str, Any]:
        """
        详细解析 PDF（返回更多信息）

        Args:
            content: PDF 文件内容
            filename: 文件名

        Returns:
            Dict: {
                "text": str,           # 合并后的文本
                "pages": List[str],    # 每页文本
                "blocks": List[Dict],  # 文本块（带坐标）
                "tables": List[Dict],  # 表格
                "metadata": Dict,      # 元数据
            }
        """
        # 保存到临时文件
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            from .pymupdf_parser import PyMuPDFParser

            parser = PyMuPDFParser(
                remove_headers=True,
                remove_footers=True,
                merge_cross_page=True,
                detect_columns=True,
            )

            result = parser.parse(tmp_path)
            logger.info(f"PDF 详细解析成功: {result['metadata']}")
            return result

        except ImportError:
            logger.warning("PyMuPDF 未安装，使用基础解析")
            text = self._parse_pdf(content, filename)
            return {
                "text": text,
                "pages": text.split("\n\n"),
                "blocks": [],
                "tables": [],
                "metadata": {"page_count": 0, "block_count": 0, "table_count": 0},
            }
        except Exception as e:
            logger.error(f"PDF 详细解析失败: {e}")
            text = self._parse_pdf(content, filename)
            return {
                "text": text,
                "pages": text.split("\n\n"),
                "blocks": [],
                "tables": [],
                "metadata": {"page_count": 0, "block_count": 0, "table_count": 0},
            }
        finally:
            # 清理临时文件
            import os
            try:
                os.unlink(tmp_path)
            except:
                pass

    def _parse_pdf(self, content: bytes, filename: str) -> str:
        """
        解析 PDF

        优先级：PyMuPDF > pdfplumber > pypdf
        """
        # 方案 1: PyMuPDF
        if self.pdf_parser == "pymupdf":
            try:
                result = self.parse_pdf_detailed(content, filename)
                return result.get("text", "")
            except Exception as e:
                logger.warning(f"PyMuPDF 解析失败: {e}")

        # 方案 2: pdfplumber
        try:
            return self._parse_pdf_pdfplumber(content)
        except Exception as e:
            logger.warning(f"pdfplumber 解析失败: {e}")

        # 方案 3: pypdf
        return self._parse_pdf_pypdf(content)

    def _parse_pdf_pdfplumber(self, content: bytes) -> str:
        """使用 pdfplumber 解析 PDF"""
        import pdfplumber

        with pdfplumber.open(io.BytesIO(content)) as pdf:
            texts = []
            for i, page in enumerate(pdf.pages):
                try:
                    text = page.extract_text()
                    if text:
                        texts.append(text)

                    # 提取表格
                    tables = page.extract_tables()
                    for table in tables:
                        table_text = self._format_table(table)
                        if table_text:
                            texts.append(table_text)
                except Exception as e:
                    logger.warning(f"PDF 第 {i+1} 页解析失败: {e}")

            if texts:
                logger.info(f"PDF 解析成功 (pdfplumber): {len(texts)} 页")
                return "\n\n".join(texts)

        raise ValueError("pdfplumber 无法提取文本")

    def _parse_pdf_pypdf(self, content: bytes) -> str:
        """使用 pypdf 解析 PDF"""
        try:
            from pypdf import PdfReader
        except ImportError:
            from PyPDF2 import PdfReader

        reader = PdfReader(io.BytesIO(content))
        texts = []
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text()
                if text:
                    texts.append(text)
            except Exception as e:
                logger.warning(f"PDF 第 {i+1} 页解析失败: {e}")

        logger.info(f"PDF 解析成功 (pypdf): {len(texts)} 页")
        return "\n\n".join(texts)

    def _parse_text(self, content: bytes) -> str:
        """解析纯文本（带编码检测）"""
        try:
            import chardet
            detected = chardet.detect(content)
            encoding = detected.get('encoding') or 'utf-8'
            confidence = detected.get('confidence', 0)

            logger.info(f"检测到编码: {encoding} (置信度: {confidence:.2%})")

            try:
                return content.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                pass
        except ImportError:
            logger.debug("chardet 未安装，使用默认编码检测")

        for encoding in ("utf-8", "gbk", "gb2312", "gb18030", "latin-1"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue

        return content.decode("utf-8", errors="ignore")

    def _parse_docx_enhanced(self, content: bytes) -> str:
        """解析 Word 文档"""
        from docx import Document

        doc = Document(io.BytesIO(content))
        texts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                texts.append(text)

        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(cells))
            if table_rows:
                texts.append("\n".join(table_rows))

        logger.info(f"Word 文档解析成功: {len(texts)} 段落/表格")
        return "\n\n".join(texts)

    def _format_table(self, table: list) -> str:
        """格式化表格为文本"""
        if not table or not table[0]:
            return ""

        rows = []
        for row in table:
            cells = [str(cell or "").strip() for cell in row]
            rows.append(" | ".join(cells))

        return "\n".join(rows)

    def get_metadata(self, filename: str, title: Optional[str] = None) -> dict:
        """获取文档元数据"""
        return {
            "filename": filename,
            "title": title or Path(filename).stem,
            "extension": Path(filename).suffix.lower(),
        }
